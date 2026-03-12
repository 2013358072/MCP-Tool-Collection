from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from fastmcp import FastMCP

from mcp_toolkit.core import config as _cfg
from mcp_toolkit.providers.base import BaseProvider

# ======================================================================
# 本模块仅处理 .docx 格式文档（python-docx）
# ======================================================================


def _mkparent(p: Path) -> None:
    p.parent.mkdir(parents=True, exist_ok=True)


def _resolve_doc_path(path: str) -> Path:
    """将路径限制在文档沙箱目录内。"""
    sandbox_root = Path(_cfg.FILESYSTEM_ROOT or Path.cwd()).resolve()
    p = Path(path)
    resolved = (sandbox_root / p).resolve() if not p.is_absolute() else p.resolve()
    try:
        resolved.relative_to(sandbox_root)
    except ValueError as e:
        raise PermissionError(
            f"路径越出沙箱: {resolved}（FILESYSTEM_ROOT={sandbox_root}）"
        ) from e
    return resolved


def _check_docx(p: Path) -> Optional[Dict[str, Any]]:
    """校验文件存在且为 .docx，不满足时返回错误字典，否则返回 None。"""
    if not p.exists() or not p.is_file():
        return {"ok": False, "error": "NOT_FOUND", "path": str(p)}
    if p.suffix.lower() != ".docx":
        return {"ok": False, "error": "NOT_DOCX", "detail": f"仅支持 .docx 格式，收到: {p.suffix}"}
    return None


def _extract_tables(doc: Document) -> List[Dict[str, Any]]:
    """提取文档中的表格内容。"""
    tables: List[Dict[str, Any]] = []
    for table_index, table in enumerate(doc.tables):
        rows: List[List[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(
            {
                "index": table_index,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
                "rows": rows,
            }
        )
    return tables


def _full_text(doc: Document) -> str:
    """提取文档全文，包含段落和表格文本。"""
    parts: List[str] = [para.text for para in doc.paragraphs if para.text]
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


# ---------------------------------------------------------------------- #
# docs_read
# ---------------------------------------------------------------------- #

def _docs_read(path: str) -> Dict[str, Any]:
    try:
        p = _resolve_doc_path(path)
    except PermissionError as e:
        return {"ok": False, "error": "PATH_OUTSIDE_SANDBOX", "detail": str(e)}
    if err := _check_docx(p):
        return err
    try:
        doc = Document(p)
        tables = _extract_tables(doc)
        paragraphs = [
            {"index": i, "text": para.text, "style": para.style.name}
            for i, para in enumerate(doc.paragraphs)
        ]
        return {
            "ok": True,
            "path": str(p),
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "content": _full_text(doc),
            "paragraphs": paragraphs,
            "tables": tables,
            "size_bytes": p.stat().st_size,
        }
    except Exception as e:
        return {"ok": False, "error": "READ_FAILED", "path": str(p), "detail": str(e)}


# ---------------------------------------------------------------------- #
# docs_write
# ---------------------------------------------------------------------- #

def _docs_write(path: str, content: str) -> Dict[str, Any]:
    """整体覆盖写入，每行作为一个独立段落。"""
    try:
        p = _resolve_doc_path(path)
    except PermissionError as e:
        return {"ok": False, "error": "PATH_OUTSIDE_SANDBOX", "detail": str(e)}
    if p.suffix.lower() != ".docx":
        return {"ok": False, "error": "NOT_DOCX", "detail": f"仅支持 .docx 格式，收到: {p.suffix}"}
    try:
        _mkparent(p)
        doc = Document()
        for line in content.splitlines():
            doc.add_paragraph(line)
        doc.save(p)
        return {"ok": True, "path": str(p), "size_bytes": p.stat().st_size}
    except Exception as e:
        return {"ok": False, "error": "WRITE_FAILED", "path": str(p), "detail": str(e)}


# ---------------------------------------------------------------------- #
# docs_append
# ---------------------------------------------------------------------- #

def _docs_append(path: str, content: str) -> Dict[str, Any]:
    """在文档末尾追加段落，每行作为一个独立段落。"""
    try:
        p = _resolve_doc_path(path)
    except PermissionError as e:
        return {"ok": False, "error": "PATH_OUTSIDE_SANDBOX", "detail": str(e)}
    if err := _check_docx(p):
        return err
    try:
        doc = Document(p)
        for line in content.splitlines():
            doc.add_paragraph(line)
        doc.save(p)
        return {
            "ok": True,
            "path": str(p),
            "paragraph_count": len(doc.paragraphs),
            "size_bytes": p.stat().st_size,
        }
    except Exception as e:
        return {"ok": False, "error": "APPEND_FAILED", "path": str(p), "detail": str(e)}


# ---------------------------------------------------------------------- #
# docs_replace
# ---------------------------------------------------------------------- #

def _docs_replace(
    path: str,
    find: str,
    replace: str,
    max_replacements: Optional[int] = None,
) -> Dict[str, Any]:
    """在文档各段落的 run 中查找并替换文本，保留原有样式。"""
    try:
        p = _resolve_doc_path(path)
    except PermissionError as e:
        return {"ok": False, "error": "PATH_OUTSIDE_SANDBOX", "detail": str(e)}
    if err := _check_docx(p):
        return err
    try:
        doc = Document(p)
        count = 0
        for para in doc.paragraphs:
            if find not in para.text:
                continue
            for run in para.runs:
                if find not in run.text:
                    continue
                occurrences = run.text.count(find)
                if max_replacements is not None:
                    allowed = max_replacements - count
                    if allowed <= 0:
                        break
                    run.text = run.text.replace(find, replace, allowed)
                    count += min(occurrences, allowed)
                else:
                    run.text = run.text.replace(find, replace)
                    count += occurrences
        doc.save(p)
        return {"ok": True, "path": str(p), "replacements": count}
    except Exception as e:
        return {"ok": False, "error": "REPLACE_FAILED", "path": str(p), "detail": str(e)}


# ---------------------------------------------------------------------- #
# docs_export_docx
# ---------------------------------------------------------------------- #

# 段落样式名映射（模型传入简写 → python-docx 样式名）
_STYLE_MAP: Dict[str, str] = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "title": "Title",
    "normal": "Normal",
    "body": "Normal",
    "quote": "Quote",
    "code": "No Spacing",
}


def _docs_export_docx(
    path: str,
    paragraphs: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """根据段落列表构建 DOCX 文档并导出。

    paragraphs 中每项格式：
        {
            "text":  str,                        # 段落文本（必填）
            "style": str,                        # 样式，支持 h1/h2/h3/h4/title/normal/quote/code
                                                 # 或 python-docx 原生样式名，默认 "Normal"
            "bold":  bool,                       # 是否加粗（可选）
            "runs":  [{"text": str, "bold": bool, "italic": bool, "underline": bool}]
                                                 # 精细 run 控制（与 bold 互斥，优先 runs）
        }
    """
    try:
        p = _resolve_doc_path(path)
    except PermissionError as e:
        return {"ok": False, "error": "PATH_OUTSIDE_SANDBOX", "detail": str(e)}
    if p.suffix.lower() != ".docx":
        return {"ok": False, "error": "NOT_DOCX", "detail": f"输出路径必须以 .docx 结尾，收到: {p.suffix}"}
    try:
        _mkparent(p)
        doc = Document()

        for item in paragraphs:
            text: str = item.get("text", "")
            raw_style: str = str(item.get("style", "normal")).lower()
            style_name: str = _STYLE_MAP.get(raw_style, item.get("style", "Normal"))

            runs: Optional[List[Dict[str, Any]]] = item.get("runs")

            if runs:
                # 精细 run 控制：先建段落，再逐 run 添加
                para = doc.add_paragraph(style=style_name)
                for run_item in runs:
                    run = para.add_run(run_item.get("text", ""))
                    if run_item.get("bold"):
                        run.bold = True
                    if run_item.get("italic"):
                        run.italic = True
                    if run_item.get("underline"):
                        run.underline = True
            else:
                para = doc.add_paragraph(text, style=style_name)
                if item.get("bold"):
                    for run in para.runs:
                        run.bold = True

        doc.save(p)
        return {
            "ok": True,
            "path": str(p),
            "paragraph_count": len(doc.paragraphs),
            "size_bytes": p.stat().st_size,
        }
    except Exception as e:
        return {"ok": False, "error": "EXPORT_DOCX_FAILED", "path": str(p), "detail": str(e)}


# ---------------------------------------------------------------------- #
# docs_find
# ---------------------------------------------------------------------- #

def _docs_find(
    path: str,
    keyword: str,
    context_chars: int = 80,
) -> Dict[str, Any]:
    """在文档全文中查找关键字，返回命中的位置、行号和上下文片段。"""
    try:
        p = _resolve_doc_path(path)
    except PermissionError as e:
        return {"ok": False, "error": "PATH_OUTSIDE_SANDBOX", "detail": str(e)}
    if err := _check_docx(p):
        return err
    try:
        doc = Document(p)
        content = _full_text(doc)
        hits: List[Dict[str, Any]] = []
        start = 0
        while True:
            idx = content.find(keyword, start)
            if idx == -1:
                break
            s = max(0, idx - context_chars)
            e = min(len(content), idx + len(keyword) + context_chars)
            hits.append({
                "position": idx,
                "line": content[:idx].count("\n") + 1,
                "snippet": content[s:e].replace("\n", "\\n"),
            })
            start = idx + len(keyword)
        return {
            "ok": True,
            "path": str(p),
            "keyword": keyword,
            "count": len(hits),
            "hits": hits,
        }
    except Exception as e:
        return {"ok": False, "error": "FIND_FAILED", "path": str(p), "detail": str(e)}


# ======================================================================
# Provider
# ======================================================================

class DocsProvider(BaseProvider):
    """DOCX 文档操作工具集 Provider（仅处理 .docx 格式）。"""

    @property
    def name(self) -> str:
        return "docs"

    def is_available(self) -> bool:
        return True

    async def initialize(self) -> None:
        self.logger.local("info", "DocsProvider 初始化完成")

    def register(self, mcp: FastMCP) -> None:

        @mcp.tool()
        async def docs_read(path: str) -> Dict[str, Any]:
            """读取 DOCX 文档，返回全文内容及段落列表（含段落序号和样式名）"""
            return _docs_read(path)

        @mcp.tool()
        async def docs_write(path: str, content: str) -> Dict[str, Any]:
            """整体覆盖写入 DOCX 文档，每行作为一个独立段落；文件不存在时自动创建"""
            return _docs_write(path, content)

        @mcp.tool()
        async def docs_append(path: str, content: str) -> Dict[str, Any]:
            """在 DOCX 文档末尾追加段落，每行作为一个独立段落"""
            return _docs_append(path, content)

        @mcp.tool()
        async def docs_replace(
            path: str,
            find: str,
            replace: str,
            max_replacements: Optional[int] = None,
        ) -> Dict[str, Any]:
            """在 DOCX 文档中按字符串匹配替换文本，保留原有字体样式；max_replacements 为空时替换全部"""
            return _docs_replace(path, find, replace, max_replacements=max_replacements)

        @mcp.tool()
        async def docs_find(
            path: str,
            keyword: str,
            context_chars: int = 80,
        ) -> Dict[str, Any]:
            """在 DOCX 文档中查找关键字，返回所有命中的位置、行号和上下文片段"""
            return _docs_find(path, keyword, context_chars=context_chars)

        @mcp.tool()
        async def docs_export_docx(
            path: str,
            paragraphs: List[Dict[str, Any]],
        ) -> Dict[str, Any]:
            """根据段落列表构建并导出 DOCX 文档。

            paragraphs 每项字段：
            - text (str, 必填): 段落文本
            - style (str, 可选): h1/h2/h3/h4/title/normal/quote/code，默认 normal
            - bold (bool, 可选): 整段加粗
            - runs (list, 可选): 精细 run 控制，每项含 text/bold/italic/underline
            """
            return _docs_export_docx(path, paragraphs)
